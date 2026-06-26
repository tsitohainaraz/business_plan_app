"""
╔══════════════════════════════════════════════════════════════╗
║  CRÉATEUR DE BUSINESS PLAN — VERSION MADAGASCAR 🇲🇬          ║
║  Powered by Tsitohaina Razafindrajoa  |  Ariary (Ar)  |  Format AFD/BNI    ║
╚══════════════════════════════════════════════════════════════╝
"""

import streamlit as st
import io
import os
import json
import requests
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.chart import BarChart, Reference, LineChart
from openpyxl.chart import PieChart

# ─── Page config ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Plan Buisness Madagascar 🇲🇬",
    page_icon="🇲🇬",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ─── COULEURS MADAGASCAR ─────────────────────────────────────────────────────
# Rouge et vert du drapeau malgache + or entrepreneurial
C_RED    = "#C8102E"   # Rouge Mada
C_GREEN  = "#007A3D"   # Vert Mada
C_DARK   = "#1A1612"   # Noir chaud
C_GOLD   = "#D4A017"   # Or entrepreneur
C_LIGHT  = "#FDF8F0"   # Fond crème
C_CARD   = "#FFFFFF"
C_BLUE   = "#1E3A5F"   # Bleu foncé pro

# ─── CSS MAGNIFIQUE ──────────────────────────────────────────────────────────
st.markdown(f"""
<style>
  @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;500;600;700;800&family=Playfair+Display:wght@700&display=swap');

  html, body, [class*="css"] {{
    font-family: 'Poppins', sans-serif;
  }}
  .main {{
    background: linear-gradient(160deg, #FDF8F0 0%, #F0F4FF 100%);
  }}

  /* ── HERO BANNER ── */
  .hero-banner {{
    background: linear-gradient(135deg, {C_DARK} 0%, {C_RED} 45%, {C_GREEN} 100%);
    border-radius: 24px;
    padding: 48px 40px 44px;
    text-align: center;
    color: white;
    margin-bottom: 28px;
    position: relative;
    overflow: hidden;
    box-shadow: 0 20px 60px rgba(200,16,46,0.25);
  }}
  .hero-banner::before {{
    content: '🇲🇬';
    position: absolute;
    font-size: 180px;
    opacity: 0.05;
    top: -30px; left: -20px;
  }}
  .hero-banner::after {{
    content: '🚀';
    position: absolute;
    font-size: 120px;
    opacity: 0.06;
    bottom: -20px; right: 20px;
  }}
  .hero-title {{
    font-family: 'Playfair Display', serif;
    font-size: 2.8em;
    font-weight: 700;
    margin-bottom: 8px;
    text-shadow: 0 2px 10px rgba(0,0,0,0.3);
    letter-spacing: 1px;
  }}
  .hero-sub {{
    font-size: 1.05em;
    opacity: 0.92;
    font-weight: 300;
    letter-spacing: 0.5px;
  }}
  .hero-badge {{
    display: inline-block;
    background: rgba(255,255,255,0.15);
    border: 1px solid rgba(255,255,255,0.3);
    border-radius: 20px;
    padding: 4px 16px;
    font-size: 0.82em;
    margin-top: 12px;
    backdrop-filter: blur(10px);
  }}

  /* ── STEP CARD ── */
  .step-card {{
    background: white;
    border-radius: 20px;
    padding: 28px 32px;
    box-shadow: 0 6px 30px rgba(0,0,0,0.07);
    border-top: 4px solid {C_RED};
    margin-bottom: 24px;
    position: relative;
  }}
  .step-card h2 {{
    color: {C_DARK};
    font-size: 1.45em;
    font-weight: 700;
    margin-bottom: 4px;
  }}
  .step-card .sub {{
    color: #7C8DB5;
    font-size: 0.9em;
    margin-bottom: 0;
  }}
  .step-number {{
    position: absolute;
    top: 24px; right: 28px;
    background: linear-gradient(135deg, {C_RED}, {C_GREEN});
    color: white;
    width: 38px; height: 38px;
    border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    font-weight: 700; font-size: 0.95em;
  }}

  /* ── PROGRESS ── */
  .prog-wrap {{
    background: #E8ECF5;
    border-radius: 50px;
    height: 12px;
    margin-bottom: 20px;
    overflow: hidden;
    box-shadow: inset 0 2px 4px rgba(0,0,0,0.08);
  }}
  .prog-fill {{
    background: linear-gradient(90deg, {C_RED}, {C_GOLD}, {C_GREEN});
    height: 100%;
    border-radius: 50px;
    transition: width 0.5s cubic-bezier(0.4, 0, 0.2, 1);
  }}
  .step-indicator {{
    display: flex;
    justify-content: center;
    gap: 8px;
    margin-bottom: 28px;
    flex-wrap: wrap;
  }}
  .step-dot {{
    width: 34px; height: 34px;
    border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    font-size: 0.78em; font-weight: 600;
    border: 2px solid #D1D9E8;
    color: #A0ABC0;
    transition: all 0.3s;
    cursor: default;
  }}
  .step-dot.active {{
    background: linear-gradient(135deg, {C_RED}, {C_DARK});
    color: white;
    border-color: {C_RED};
    box-shadow: 0 4px 12px rgba(200,16,46,0.35);
    transform: scale(1.12);
  }}
  .step-dot.done {{
    background: {C_GREEN};
    color: white;
    border-color: {C_GREEN};
  }}

  /* ── INFOBOX ── */
  .info-box {{
    background: linear-gradient(135deg, #FFF8E1, #FFFDE7);
    border-left: 4px solid {C_GOLD};
    border-radius: 10px;
    padding: 12px 16px;
    font-size: 0.87em;
    color: #5D4037;
    margin-bottom: 18px;
  }}
  .ai-box {{
    background: linear-gradient(135deg, #E8F5E9, #F3E5F5);
    border-left: 4px solid {C_GREEN};
    border-radius: 10px;
    padding: 14px 16px;
    font-size: 0.88em;
    color: #1B5E20;
    margin-top: 10px;
    margin-bottom: 12px;
  }}

  /* ── FEATURE CARDS (page accueil) ── */
  .feat-card {{
    background: white;
    border-radius: 16px;
    padding: 22px 18px;
    text-align: center;
    box-shadow: 0 4px 20px rgba(0,0,0,0.06);
    border-bottom: 3px solid transparent;
    transition: all 0.3s;
  }}
  .feat-card:hover {{ transform: translateY(-4px); }}
  .feat-card.red {{ border-bottom-color: {C_RED}; }}
  .feat-card.green {{ border-bottom-color: {C_GREEN}; }}
  .feat-card.gold {{ border-bottom-color: {C_GOLD}; }}
  .feat-card.blue {{ border-bottom-color: {C_BLUE}; }}
  .feat-icon {{ font-size: 2.4em; margin-bottom: 8px; }}
  .feat-title {{ font-weight: 700; font-size: 1em; color: {C_DARK}; }}
  .feat-desc {{ font-size: 0.82em; color: #8A94A6; margin-top: 4px; }}

  /* ── SUCCESS BOX ── */
  .success-box {{
    background: linear-gradient(135deg, #F0FFF4, #E8F5E9);
    border: 2px solid {C_GREEN};
    border-radius: 20px;
    padding: 36px 32px;
    text-align: center;
    margin-bottom: 24px;
    box-shadow: 0 8px 32px rgba(0,122,61,0.12);
  }}
  .success-box h2 {{ color: {C_GREEN}; font-size: 2em; font-weight: 800; }}
  .success-box p  {{ color: #2E7D32; font-size: 1.05em; }}

  /* ── BUTTONS ── */
  div[data-testid="stButton"] > button {{
    border-radius: 12px !important;
    font-weight: 600 !important;
    padding: 10px 26px !important;
    transition: all .25s !important;
    font-family: 'Poppins', sans-serif !important;
  }}
  div[data-testid="stButton"] > button:first-child {{
    background: linear-gradient(135deg, {C_RED}, #A0001A) !important;
    color: white !important;
    border: none !important;
    box-shadow: 0 4px 15px rgba(200,16,46,0.3) !important;
  }}
  div[data-testid="stButton"] > button:first-child:hover {{
    transform: translateY(-2px);
    box-shadow: 0 8px 25px rgba(200,16,46,0.45) !important;
  }}
  div[data-testid="stDownloadButton"] > button {{
    border-radius: 12px !important;
    font-weight: 700 !important;
    padding: 13px 20px !important;
    width: 100% !important;
    font-size: 0.95em !important;
    font-family: 'Poppins', sans-serif !important;
  }}

  /* ── SECTION TITLE ── */
  .section-title {{
    font-family: 'Playfair Display', serif;
    font-size: 1.1em;
    color: {C_RED};
    font-weight: 700;
    border-bottom: 2px solid #FFE4E8;
    padding-bottom: 6px;
    margin: 18px 0 12px;
  }}

  /* ── FINANCIER CARD ── */
  .fin-card {{
    background: linear-gradient(135deg, {C_DARK} 0%, {C_BLUE} 100%);
    color: white;
    border-radius: 14px;
    padding: 18px 20px;
    text-align: center;
    box-shadow: 0 6px 20px rgba(30,58,95,0.3);
  }}
  .fin-card .amount {{ font-size: 1.6em; font-weight: 800; color: {C_GOLD}; }}
  .fin-card .label  {{ font-size: 0.82em; opacity: 0.85; margin-top: 4px; }}

  /* ── TICKER animé ── */
  @keyframes slideIn {{
    from {{ opacity: 0; transform: translateX(-20px); }}
    to   {{ opacity: 1; transform: translateX(0); }}
  }}
  .animate-in {{ animation: slideIn 0.4s ease forwards; }}

  /* ── BADGE SECTEUR ── */
  .sector-badge {{
    display: inline-block;
    background: linear-gradient(135deg, {C_GREEN}, #00A651);
    color: white;
    border-radius: 20px;
    padding: 3px 12px;
    font-size: 0.78em;
    font-weight: 600;
    margin-left: 8px;
  }}
</style>
""", unsafe_allow_html=True)

# ─── Session state ────────────────────────────────────────────────────────────
if "step" not in st.session_state:
    st.session_state.step = 0
if "data" not in st.session_state:
    st.session_state.data = {}
if "ai_suggestions" not in st.session_state:
    st.session_state.ai_suggestions = {}

STEPS = [
    "🏠 Accueil", "💡 Votre Idée", "🏢 Votre Projet", "📍 Localisation",
    "📈 Marché", "📦 Produit/Service", "👥 Équipe",
    "💰 Financement", "📊 Projections", "🌱 Impact Social",
    "⚖️ Statut Juridique", "🎉 Récapitulatif"
]
TOTAL_STEPS = len(STEPS)

# ─── HELPERS ─────────────────────────────────────────────────────────────────
def next_step(): st.session_state.step += 1
def prev_step(): st.session_state.step -= 1

def fmt_ariary(val):
    """Formate un nombre en Ariary."""
    try:
        v = int(val)
        return f"{v:,} Ar".replace(",", " ")
    except:
        return "0 Ar"

def nav_buttons(show_prev=True):
    c1, c2, c3 = st.columns([1, 3, 1])
    if show_prev:
        with c1:
            if st.button("◀ Précédent", key=f"prev_{st.session_state.step}"):
                prev_step()
                st.rerun()
    with c3:
        if st.button("Suivant ▶", key=f"next_{st.session_state.step}"):
            next_step()
            st.rerun()

def render_progress():
    step = st.session_state.step
    pct = int((step / (TOTAL_STEPS - 1)) * 100)
    st.markdown(f"""
    <div class="prog-wrap">
      <div class="prog-fill" style="width:{pct}%"></div>
    </div>""", unsafe_allow_html=True)

    dots = ""
    for i, name in enumerate(STEPS):
        emoji = name.split()[0]
        if i < step:
            cls, icon = "done", "✓"
        elif i == step:
            cls, icon = "active", emoji
        else:
            cls, icon = "", str(i + 1)
        dots += f'<div class="step-dot {cls}" title="{name}">{icon}</div>'
    st.markdown(f'<div class="step-indicator">{dots}</div>', unsafe_allow_html=True)

def field(label, key, placeholder="", multiline=False, help=""):
    val = st.session_state.data.get(key, "")
    if multiline:
        result = st.text_area(label, value=val, placeholder=placeholder, help=help, height=110)
    else:
        result = st.text_input(label, value=val, placeholder=placeholder, help=help)
    st.session_state.data[key] = result
    return result

def num_field(label, key, min_val=0, max_val=500_000_000_000, step=100_000, default=0, fmt="%d", help=""):
    val = st.session_state.data.get(key, default)
    result = st.number_input(label, min_value=min_val, max_value=max_val,
                             value=int(val), step=step, format=fmt, help=help)
    st.session_state.data[key] = result
    return result

def select_field(label, key, options, help=""):
    val = st.session_state.data.get(key, options[0])
    idx = options.index(val) if val in options else 0
    result = st.selectbox(label, options, index=idx, help=help)
    st.session_state.data[key] = result
    return result

def radio_field(label, key, options):
    val = st.session_state.data.get(key, options[0])
    idx = options.index(val) if val in options else 0
    result = st.radio(label, options, index=idx, horizontal=True)
    st.session_state.data[key] = result
    return result

def info(txt):
    st.markdown(f'<div class="info-box">💡 {txt}</div>', unsafe_allow_html=True)

def section_title(txt):
    st.markdown(f'<div class="section-title">{txt}</div>', unsafe_allow_html=True)

# ─── CLAUDE AI SUGGESTIONS ────────────────────────────────────────────────────
def get_ai_suggestion(section_key, context_data, prompt_template):
    """Appel Claude API pour générer une suggestion."""
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return None

    context_str = "\n".join([f"- {k}: {v}" for k, v in context_data.items() if v])
    prompt = f"""{prompt_template}

Contexte du projet à Madagascar:
{context_str}

Réponds en français, de manière concise et pratique (3-5 phrases maximum).
Contexte: Madagascar, marché local, monnaie = Ariary (Ar), public = entrepreneurs malgaches."""

    try:
        headers = {
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01"
        }
        payload = {
            "model": "claude-sonnet-4-6",
            "max_tokens": 400,
            "messages": [{"role": "user", "content": prompt}]
        }
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers=headers, json=payload, timeout=15
        )
        if resp.status_code == 200:
            return resp.json()["content"][0]["text"]
    except Exception:
        pass
    return None

def ai_suggest_button(section_key, label, context_data, prompt_template):
    """Bouton IA avec résultat stocké."""
    col_btn, col_tip = st.columns([1, 3])
    with col_btn:
        if st.button(f"🤖 Suggestion IA", key=f"ai_{section_key}"):
            with st.spinner("Génération en cours…"):
                suggestion = get_ai_suggestion(section_key, context_data, prompt_template)
                if suggestion:
                    st.session_state.ai_suggestions[section_key] = suggestion
                else:
                    st.session_state.ai_suggestions[section_key] = (
                        "⚠️ Clé API non configurée. Ajoutez ANTHROPIC_API_KEY dans les variables d'environnement Streamlit."
                    )
    with col_tip:
        st.caption("🇲🇬 Conseil adapté Madagascar par Claude AI")

    if section_key in st.session_state.ai_suggestions:
        st.markdown(f'<div class="ai-box">🤖 <strong>Suggestion IA :</strong><br>{st.session_state.ai_suggestions[section_key]}</div>',
                    unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════
# ÉTAPES
# ══════════════════════════════════════════════════════════════

def step_0_welcome():
    st.markdown("""
    <div class="hero-banner">
      <div class="hero-title">🚀 Créateur de Business Plan</div>
      <div class="hero-sub">Votre consultant en affaires numérique — Spécialisé Madagascar 🇲🇬</div>
      <div class="hero-badge">✨ Format AFD · BNI · HABAKA · Bailleurs Internationaux</div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3, col4 = st.columns(4)
    cards = [
        ("red", "📋", "11 Étapes", "Questions guidées simples"),
        ("green", "🇲🇬", "100% Madagascar", "Ariary, régions, bailleurs locaux"),
        ("gold", "🤖", "Propulsé par l'IA", "Suggestions Claude sur chaque section"),
        ("blue", "📄📊", "2 Documents Pro", "Word complet + Excel financier"),
    ]
    for col, (color, icon, title, desc) in zip([col1, col2, col3, col4], cards):
        with col:
            st.markdown(f"""
            <div class="feat-card {color}">
              <div class="feat-icon">{icon}</div>
              <div class="feat-title">{title}</div>
              <div class="feat-desc">{desc}</div>
            </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    st.markdown("""
    <div style="background:white;border-radius:16px;padding:24px 28px;box-shadow:0 4px 20px rgba(0,0,0,0.06);margin-bottom:20px;">
    <strong>📌 Ce que vous obtiendrez :</strong><br><br>
    ✅ Business Plan Word professionnel (25+ pages) — format reconnu par BNI, BOA, AFD, HABAKA<br>
    ✅ Plan Financier Excel sur 5 ans — en Ariary — avec Cash Flow, SWOT, Bilan, Point mort<br>
    ✅ Section Impact Social obligatoire pour les bailleurs internationaux<br>
    ✅ Analyse des risques pays Madagascar<br>
    ✅ Conseils IA personnalisés à votre secteur d'activité
    </div>
    """, unsafe_allow_html=True)

    _, c, _ = st.columns([2, 1, 2])
    with c:
        if st.button("🚀 Commencer maintenant", key="start"):
            next_step()
            st.rerun()


def step_1_idea():
    st.markdown(f"""<div class="step-card animate-in">
    <div class="step-number">1</div>
    <h2>💡 Votre Idée</h2>
    <div class="sub">Le problème que vous résolvez et votre solution pour le marché malgache.</div>
    </div>""", unsafe_allow_html=True)

    info("Une bonne idée répond à un besoin réel de la population malgache. Pensez local d'abord !")

    field("Quel problème ou besoin résolvez-vous à Madagascar ?", "probleme",
          "Ex : Les agriculteurs de la région Boeny n'ont pas accès à des semences de qualité…", multiline=True)
    field("Quelle est votre solution concrète ?", "solution",
          "Ex : Une plateforme de mise en relation entre pépiniéristes et agriculteurs…", multiline=True)

    col1, col2 = st.columns(2)
    with col1:
        select_field("Secteur d'activité", "domaine", [
            "Agriculture / Élevage / Pêche",
            "Commerce / Négoce / Import-Export",
            "Services aux entreprises (B2B)",
            "Services aux particuliers (B2C)",
            "Restauration / Hôtellerie / Tourisme",
            "Artisanat / Textile / Mode",
            "Santé / Pharmacie / Bien-être",
            "Éducation / Formation / EdTech",
            "Énergie / Environnement / Eau",
            "Tech / Numérique / Mobile Money",
            "BTP / Immobilier / Construction",
            "Transport / Logistique",
            "Industrie / Transformation",
            "Autre"
        ])
    with col2:
        radio_field("Votre cible principale ?", "cible",
                    ["B2C (Particuliers)", "B2B (Entreprises)", "Les deux"])

    field("Concurrents existants à Madagascar ou dans la région ?", "concurrents",
          "Ex : Marchés locaux, importateurs…", multiline=True)
    field("Votre avantage compétitif face à eux ?", "differenciateur",
          "Ex : Prix plus bas, qualité supérieure, service après-vente…", multiline=True)

    ai_suggest_button(
        "idea_ai",
        "Suggestion IA",
        {"probleme": st.session_state.data.get("probleme",""),
         "solution": st.session_state.data.get("solution",""),
         "domaine": st.session_state.data.get("domaine","")},
        "Donne-moi 3 conseils pour affiner cette idée d'entreprise à Madagascar et la rendre plus attractive pour les investisseurs et bailleurs locaux."
    )

    nav_buttons(show_prev=False)


def step_2_project():
    st.markdown(f"""<div class="step-card animate-in">
    <div class="step-number">2</div>
    <h2>🏢 Votre Projet</h2>
    <div class="sub">Présentez votre projet et vos motivations personnelles.</div>
    </div>""", unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        field("Nom de l'entreprise / projet", "nom_entreprise", "Ex : AgroMada SARL")
        field("Prénom et nom du porteur de projet", "porteur", "Ex : Rakoto Jean")
    with col2:
        field("Date de création prévue", "date_creation", "Ex : Janvier 2027")
        radio_field("Vous lancez-vous seul ou à plusieurs ?", "seul_ou_plusieurs",
                    ["Seul(e)", "À 2", "À 3 et plus"])

    field("Pitch du projet (2-3 phrases)", "pitch",
          "Ex : Nous fournissons des semences certifiées aux agriculteurs du Menabe via un réseau de distributeurs locaux…", multiline=True)
    field("Pourquoi ce projet ? Votre motivation personnelle ?", "motivation",
          "Ex : Fils d'agriculteur, j'ai vu les difficultés de mon père…", multiline=True)
    field("Expériences / compétences dans ce domaine ?", "competences_porteur",
          "Ex : 5 ans dans l'agribusiness, formation ESSA Antananarivo…", multiline=True)

    nav_buttons()


def step_3_location():
    st.markdown(f"""<div class="step-card animate-in">
    <div class="step-number">3</div>
    <h2>📍 Localisation & Contexte Malgache</h2>
    <div class="sub">Précisez votre ancrage territorial — essentiel pour les bailleurs.</div>
    </div>""", unsafe_allow_html=True)

    info("Les bailleurs internationaux (AFD, Banque Mondiale, PNUD) exigent une analyse du contexte local de votre région.")

    col1, col2 = st.columns(2)
    with col1:
        select_field("Région principale d'activité", "region", [
            "Analamanga (Antananarivo)",
            "Vakinankaratra (Antsirabe)",
            "Itasy",
            "Bongolava",
            "Haute Matsiatra (Fianarantsoa)",
            "Amoron'i Mania",
            "Vatovavy (Mananjary)",
            "Fitovinany (Manakara)",
            "Atsimo-Atsinanana (Farafangana)",
            "Atsinanana (Toamasina)",
            "Analanjirofo (Fénérive)",
            "Alaotra-Mangoro (Ambatondrazaka)",
            "Boeny (Mahajanga)",
            "Sofia (Antsohihy)",
            "Betsiboka",
            "Melaky",
            "Atsimo-Andrefana (Toliara)",
            "Androy (Ambovombe)",
            "Anosy (Fort-Dauphin)",
            "Menabe (Morondava)",
            "Diana (Antsiranana)",
            "Sava (Sambava)",
        ])
        field("Ville / Fokontany d'implantation", "ville", "Ex : Antananarivo — Andraharo")
    with col2:
        select_field("Zone d'activité principale", "zone_activite", [
            "Urbaine (grande ville)",
            "Périurbaine",
            "Rurale",
            "Multi-régionale (plusieurs zones)",
            "Nationale",
            "Export / International"
        ])
        select_field("Rayon géographique des clients", "rayon_clients", [
            "Quartier / Commune",
            "Ville / District",
            "Région",
            "Plusieurs régions",
            "Madagascar entier",
            "Export vers l'Afrique / International"
        ])

    field("Infrastructures disponibles dans votre zone", "infrastructures",
          "Ex : Route nationale, électricité JIRAMA, connexion internet, marché hebdomadaire…", multiline=True)
    field("Contraintes ou opportunités spécifiques à votre région", "contraintes_region",
          "Ex : Cyclones fréquents dans l'Atsinanana, fort trafic touristique à Nosy-Be…", multiline=True)

    nav_buttons()


def step_4_market():
    st.markdown(f"""<div class="step-card animate-in">
    <div class="step-number">4</div>
    <h2>📈 Étude de Marché — Madagascar</h2>
    <div class="sub">Prouvez qu'il existe une vraie opportunité dans votre marché local.</div>
    </div>""", unsafe_allow_html=True)

    info("Utilisez des données locales : INSTAT Madagascar, EDBM, BNI, Chambre de Commerce.")

    col1, col2 = st.columns(2)
    with col1:
        field("Taille estimée de votre marché à Madagascar", "taille_marche",
              "Ex : 2 millions de ménages ruraux, marché de 500 Mds Ar/an")
        field("Source de cette estimation", "source_marche",
              "Ex : INSTAT 2024, rapport BNI, étude EDBM…")
    with col2:
        select_field("Maturité du marché", "maturite_marche",
                     ["Inexploité / Nouveau", "Émergent", "En croissance", "Mature / Établi"])
        field("Nb de clients potentiels visés — An 1", "nb_clients_an1",
              "Ex : 200 agriculteurs dans la région Boeny")

    field("Tendances du marché à Madagascar", "tendances",
          "Ex : Croissance de la classe moyenne à Tana, boom du mobile money, digitalisation post-COVID…", multiline=True)
    field("3 concurrents principaux (locaux ou internationaux)", "concurrents_details",
          "Nom — Localisation — Prix — Forces / Faiblesses…", multiline=True)
    field("Stratégie d'acquisition de vos premiers clients", "strategie_acquisition",
          "Ex : Bouche-à-oreille, marchés locaux, radio communautaire, mobile money promotions…", multiline=True)

    ai_suggest_button(
        "market_ai",
        "Analyse de marché IA",
        {"domaine": st.session_state.data.get("domaine",""),
         "region": st.session_state.data.get("region",""),
         "taille_marche": st.session_state.data.get("taille_marche","")},
        "Analyse le potentiel de ce marché à Madagascar pour ce secteur, et propose 3 opportunités concrètes à exploiter."
    )

    nav_buttons()


def step_5_product():
    st.markdown(f"""<div class="step-card animate-in">
    <div class="step-number">5</div>
    <h2>📦 Produit / Service</h2>
    <div class="sub">Décrivez ce que vous vendez, comment et à quel prix en Ariary.</div>
    </div>""", unsafe_allow_html=True)

    field("Description détaillée de votre produit ou service", "produit_detail",
          "Caractéristiques, fonctionnalités, avantages…", multiline=True)
    field("Comment est-il produit ou livré ?", "production",
          "Ex : Transformé dans notre atelier à Antsirabe, livré via Moov…", multiline=True)

    col1, col2 = st.columns(2)
    with col1:
        select_field("Modèle économique principal", "modele_eco", [
            "Vente directe",
            "Abonnement mensuel / annuel",
            "Commission / % sur transaction",
            "Freemium",
            "Service à la demande",
            "Marketplace / Place de marché",
            "B2B — Contrats entreprises",
            "Franchise / Licence",
            "Autre"
        ])
        num_field("Prix de vente unitaire (Ar)", "prix_vente_ar",
                  step=500, max_val=1_000_000_000,
                  help="Prix en Ariary pour une unité ou un mois")
    with col2:
        num_field("Coût de revient par unité (Ar)", "cout_revient_ar",
                  step=500, max_val=500_000_000,
                  help="Coût variable par unité vendue en Ariary")
        num_field("Marge brute estimée (%)", "marge_brute",
                  min_val=0, max_val=100, step=1, default=40,
                  help="(Prix - Coût) / Prix × 100")

    field("Ressources clés nécessaires", "ressources_cles",
          "Ex : Terrain, machine de transformation, camion, connexion internet…", multiline=True)
    field("Fournisseurs principaux à Madagascar", "fournisseurs",
          "Ex : STAR Madagascar, SMTP, grossistes de Petite Vitesse…")

    nav_buttons()


def step_6_team():
    st.markdown(f"""<div class="step-card animate-in">
    <div class="step-number">6</div>
    <h2>👥 L'Équipe</h2>
    <div class="sub">Les bailleurs financent d'abord des personnes. Mettez votre équipe en valeur.</div>
    </div>""", unsafe_allow_html=True)

    info("Les investisseurs à Madagascar apprécient les équipes avec ancrage local ET ouverture internationale.")

    nb_str = st.session_state.data.get("seul_ou_plusieurs", "Seul(e)")
    nb = 1 if "Seul" in nb_str else (2 if "À 2" in nb_str else 3)

    for i in range(1, nb + 1):
        st.markdown(f"**👤 Fondateur / Fondatrice {i}**")
        c1, c2 = st.columns(2)
        with c1:
            field(f"Nom complet", f"fond_{i}_nom", f"Ex : Rakoto Jean")
            field(f"Rôle", f"fond_{i}_role", "Ex : Directeur Général / Technique")
        with c2:
            field(f"Formation / Diplôme", f"fond_{i}_formation", "Ex : Licence en Agronomie, ESSA Tana")
            field(f"Expérience principale", f"fond_{i}_exp", "Ex : 5 ans en agribusiness")
        field(f"Compétences clés apportées", f"fond_{i}_skills",
              "Ex : Réseau paysan, gestion de projet, langues (malagasy, français, anglais)")
        st.markdown("---")

    field("Conseillers / Mentors / Partenaires clés", "advisors",
          "Ex : M. Rakotobe (ex-BNI), Incubateur HABAKA, Chambre de Commerce…")
    field("Recrutements prévus (postes et délais)", "recrutements",
          "Ex : 2 techniciens agricoles dès le 3e mois, 1 comptable an 2")

    nav_buttons()


def step_7_financing():
    st.markdown(f"""<div class="step-card animate-in">
    <div class="step-number">7</div>
    <h2>💰 Financement — Sources à Madagascar</h2>
    <div class="sub">Comment financer votre projet avec les ressources disponibles à Madagascar.</div>
    </div>""", unsafe_allow_html=True)

    info("À Madagascar, les banques exigent souvent des garanties réelles (titre foncier, matériel). Pensez à préparer vos dossiers de garantie.")

    section_title("💵 Fonds propres et apports")
    col1, col2 = st.columns(2)
    with col1:
        num_field("Apports personnels (Ar)", "apports_perso", step=500_000)
        num_field("Apports famille / diaspora (Ar)", "love_money", step=500_000)
    with col2:
        num_field("Capital social initial (Ar)", "capital_social", step=100_000, default=2_000_000)
        num_field("Budget de lancement total (Ar)", "budget_lancement", step=1_000_000)

    section_title("🏦 Financement bancaire et institutionnel")
    col1, col2 = st.columns(2)
    with col1:
        select_field("Banque / MFI visée", "banque_visee", [
            "BNI Madagascar",
            "BOA Madagascar",
            "BMOI",
            "BFV-SG",
            "AccèsBanque Madagascar",
            "CECAM",
            "Microcred / Baobab",
            "Aucune banque pour l'instant",
        ])
        num_field("Emprunt bancaire souhaité (Ar)", "emprunt", step=1_000_000)
    with col2:
        field("Garanties disponibles", "garanties",
              "Ex : Titre foncier à Ambohidratrimo, matériel agricole évalué à 50M Ar")
        num_field("Subventions / Aides publiques (Ar)", "subventions", step=500_000)

    section_title("🌍 Bailleurs internationaux et fonds")
    col1, col2 = st.columns(2)
    with col1:
        select_field("Bailleur international visé", "bailleur_vise", [
            "AFD (Agence Française de Développement)",
            "Banque Mondiale / IDA",
            "PNUD / ODD",
            "Union Européenne",
            "USAID",
            "GIZ (Allemagne)",
            "IFAD (Agriculture familiale)",
            "BAD (Banque Africaine de Développement)",
            "HABAKA (Incubateur startup Tana)",
            "Investissimo",
            "Aucun bailleur international",
        ])
    with col2:
        num_field("Montant bailleur attendu (Ar)", "montant_bailleur", step=1_000_000)
        field("À quoi servira ce financement ?", "usage_budget",
              "Ex : 40% équipement, 30% fonds de roulement, 20% formation, 10% divers")

    nav_buttons()


def step_8_projections():
    st.markdown(f"""<div class="step-card animate-in">
    <div class="step-number">8</div>
    <h2>📊 Projections Financières — 5 ans en Ariary</h2>
    <div class="sub">Estimez vos revenus et charges sur 5 ans (standard bailleurs internationaux).</div>
    </div>""", unsafe_allow_html=True)

    info("Les bailleurs AFD et Banque Mondiale exigent des projections sur 5 ans minimum en Ariary.")

    section_title("📈 Chiffre d'Affaires Prévisionnel (Ar)")
    c1, c2, c3, c4, c5 = st.columns(5)
    cas = []
    for i, col in enumerate([c1, c2, c3, c4, c5], 1):
        with col:
            v = num_field(f"CA An {i} (Ar)", f"ca_an{i}", step=1_000_000, max_val=100_000_000_000)
            cas.append(v)

    if any(cas):
        st.markdown("**📊 Aperçu de la croissance prévue :**")
        preview_cols = st.columns(5)
        for i, (col, v) in enumerate(zip(preview_cols, cas), 1):
            with col:
                growth = ""
                if i > 1 and cas[i-2] > 0:
                    g = ((v - cas[i-2]) / cas[i-2]) * 100
                    growth = f" (+{g:.0f}%)" if g > 0 else f" ({g:.0f}%)"
                st.markdown(f'<div class="fin-card"><div class="amount">{fmt_ariary(v)}</div><div class="label">Année {i}{growth}</div></div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    section_title("📉 Charges Mensuelles Fixes (Ar/mois)")
    c1, c2, c3 = st.columns(3)
    with c1:
        num_field("Salaires et rémunérations (Ar/mois)", "charges_salaires", step=50_000)
        num_field("Loyer / hébergement (Ar/mois)", "charges_loyer", step=50_000)
    with c2:
        num_field("Matières premières / Stock (Ar/mois)", "charges_matieres", step=100_000)
        num_field("Transport / Carburant (Ar/mois)", "charges_transport", step=50_000)
    with c3:
        num_field("Électricité / Eau / Téléphone (Ar/mois)", "charges_utilities", step=20_000)
        num_field("Autres charges fixes (Ar/mois)", "autres_charges", step=50_000)

    section_title("📈 Charges Variables (% du CA)")
    c1, c2 = st.columns(2)
    with c1:
        num_field("Coût des ventes / COGS (%)", "cogs_pct",
                  min_val=0, max_val=100, step=1, default=35)
    with c2:
        num_field("Commissions / frais divers (%)", "commission_pct",
                  min_val=0, max_val=30, step=1, default=5)

    section_title("💸 Investissements initiaux (Ar)")
    c1, c2, c3 = st.columns(3)
    with c1:
        num_field("Équipements / Matériels (Ar)", "invest_equipement", step=1_000_000)
    with c2:
        num_field("Aménagement / Construction (Ar)", "invest_amenagement", step=1_000_000)
    with c3:
        num_field("Fonds de roulement initial (Ar)", "invest_fdr", step=500_000)

    nav_buttons()


def step_9_impact():
    st.markdown(f"""<div class="step-card animate-in">
    <div class="step-number">9</div>
    <h2>🌱 Impact Social & Environnemental</h2>
    <div class="sub">Section OBLIGATOIRE pour tous les bailleurs internationaux à Madagascar.</div>
    </div>""", unsafe_allow_html=True)

    info("AFD, Banque Mondiale, PNUD et GIZ exigent une section détaillée sur l'impact social. C'est souvent un critère de sélection principal.")

    section_title("👥 Emplois et bénéficiaires")
    c1, c2 = st.columns(2)
    with c1:
        num_field("Emplois directs créés (An 1)", "emplois_directs_an1", step=1, max_val=10000)
        num_field("Emplois directs créés (An 3)", "emplois_directs_an3", step=1, max_val=10000)
    with c2:
        num_field("Emplois indirects estimés", "emplois_indirects", step=5, max_val=100000)
        num_field("Nb de femmes bénéficiaires directes", "femmes_beneficiaires", step=5, max_val=100000)

    field("Populations bénéficiaires principalement touchées", "populations_cibles",
          "Ex : Femmes rurales du Vakinankaratra, jeunes sans emploi de Toamasina, agriculteurs du Menabe…", multiline=True)

    section_title("🌍 ODD — Objectifs de Développement Durable ciblés")
    odd_options = [
        "ODD 1 — Pas de pauvreté",
        "ODD 2 — Faim zéro / Agriculture durable",
        "ODD 3 — Bonne santé et bien-être",
        "ODD 4 — Éducation de qualité",
        "ODD 5 — Égalité entre les sexes",
        "ODD 6 — Eau propre et assainissement",
        "ODD 7 — Énergie propre et d'un coût abordable",
        "ODD 8 — Travail décent et croissance économique",
        "ODD 9 — Industrie, innovation et infrastructure",
        "ODD 10 — Inégalités réduites",
        "ODD 11 — Villes et communautés durables",
        "ODD 12 — Consommation et production responsables",
        "ODD 13 — Mesures relatives à la lutte contre les changements climatiques",
        "ODD 15 — Vie terrestre / Biodiversité"
    ]
    odd_sel = st.multiselect("Sélectionnez les ODD auxquels votre projet contribue", odd_options)
    st.session_state.data["odds_selectionnes"] = ", ".join(odd_sel)

    field("Comment votre projet contribue-t-il à ces ODD ?", "contribution_odd",
          "Ex : Réduction de la déforestation par l'usage de foyers améliorés, accès à l'eau potable…", multiline=True)

    section_title("⚠️ Risques Pays et Atténuation")
    field("Risques identifiés à Madagascar (politique, climatique, infrastructures)", "risques_pays",
          "Ex : Instabilité politique, cyclones dans le nord-est, coupures d'électricité fréquentes, route enclavée…", multiline=True)
    field("Mesures d'atténuation et plan de contingence", "attenuation_risques",
          "Ex : Stock tampon de 3 mois, groupe électrogène de secours, partenariat avec ONG locale en cas de crise…", multiline=True)
    field("Impact environnemental de votre activité", "impact_env",
          "Ex : Réduction des émissions CO2 par substitution bois/solaire, gestion des déchets…", multiline=True)

    ai_suggest_button(
        "impact_ai",
        "Analyse d'impact IA",
        {"domaine": st.session_state.data.get("domaine",""),
         "region": st.session_state.data.get("region",""),
         "emplois_directs_an1": st.session_state.data.get("emplois_directs_an1","")},
        "Propose 4 indicateurs d'impact social mesurables pour ce projet à Madagascar, et explique comment les présenter aux bailleurs internationaux."
    )

    nav_buttons()


def step_10_legal():
    st.markdown(f"""<div class="step-card animate-in">
    <div class="step-number">10</div>
    <h2>⚖️ Statut Juridique — Droit Malgache</h2>
    <div class="sub">Choisissez la structure adaptée selon le droit malgache et votre type de financement.</div>
    </div>""", unsafe_allow_html=True)

    info("En droit malgache (OHADA), les formes les plus courantes sont SARL, SA, SARLU et l'Entreprise Individuelle (EI). Le guichet unique EDBM facilite les immatriculations.")

    col1, col2 = st.columns(2)
    with col1:
        select_field("Statut juridique envisagé", "statut_juridique", [
            "Entreprise Individuelle (EI)",
            "SARLU — Société à Responsabilité Limitée Unipersonnelle",
            "SARL — Société à Responsabilité Limitée",
            "SA — Société Anonyme",
            "Coopérative (OHADA)",
            "Association à but lucratif",
            "ONG / Association",
            "Je ne sais pas encore",
        ])
        num_field("Capital social (Ar)", "capital_social_ar",
                  step=100_000, default=2_000_000,
                  help="Minimum légal en droit malgache : 2 000 000 Ar pour une SARL")
    with col2:
        field("Adresse du siège social", "siege_social",
              "Ex : Lot II C 47, Antananarivo 101")
        select_field("Immatriculation via", "immatriculation_via", [
            "EDBM (Guichet unique — recommandé)",
            "Greffe du Tribunal de Commerce",
            "Expert-comptable / Cabinet juridique",
            "Non encore démarré",
        ])

    field("Régime fiscal envisagé", "regime_fiscal",
          "Ex : Régime du réel (BIC), Régime synthétique (micro-entreprise), TVA 20%…")
    field("Questions ou précisions sur le statut ?", "questions_statut",
          "Ex : J'hésite entre SARLU et SARL car je pense lever des fonds en An 2…", multiline=True)

    ai_suggest_button(
        "legal_ai",
        "Conseil juridique IA",
        {"statut_juridique": st.session_state.data.get("statut_juridique",""),
         "capital_social_ar": st.session_state.data.get("capital_social_ar",""),
         "bailleur_vise": st.session_state.data.get("bailleur_vise","")},
        "Donne un conseil sur le choix du statut juridique pour ce type de projet à Madagascar, en tenant compte du bailleur visé et des exigences OHADA."
    )

    nav_buttons()


def step_11_recap():
    d = st.session_state.data
    nom = d.get("nom_entreprise", "Votre Projet") or "Votre Projet"

    st.markdown(f"""<div class="success-box">
    <h2>🎉 Félicitations !</h2>
    <p>Votre Business Plan <strong>{nom}</strong> est prêt à être téléchargé.</p>
    <p>Documents générés selon le format <strong>AFD / BNI Madagascar / Bailleurs Internationaux</strong></p>
    </div>""", unsafe_allow_html=True)

    # Résumé des KPIs financiers
    ca1 = int(d.get("ca_an1", 0) or 0)
    ca3 = int(d.get("ca_an3", 0) or 0)
    budget = int(d.get("budget_lancement", 0) or 0)
    emplois = int(d.get("emplois_directs_an1", 0) or 0)

    if any([ca1, ca3, budget, emplois]):
        st.markdown("### 📊 Vos Indicateurs Clés")
        kpi_cols = st.columns(4)
        kpis = [
            ("CA Année 1", fmt_ariary(ca1)),
            ("CA Année 3", fmt_ariary(ca3)),
            ("Budget lancement", fmt_ariary(budget)),
            (f"Emplois An 1", f"{emplois} postes"),
        ]
        for col, (label, val) in zip(kpi_cols, kpis):
            with col:
                st.markdown(f'<div class="fin-card"><div class="amount">{val}</div><div class="label">{label}</div></div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("### 📥 Télécharger vos Documents")

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("#### 📄 Business Plan Word")
        st.caption("Document complet 25+ pages — Format AFD / BNI / HABAKA — Droit malgache")
        try:
            docx_bytes = generate_word(d)
            st.download_button(
                label="⬇️ Télécharger le Business Plan (.docx)",
                data=docx_bytes,
                file_name=f"BusinessPlan_{nom.replace(' ','_')}_Madagascar.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            st.error(f"Erreur génération Word : {e}")

    with c2:
        st.markdown("#### 📊 Plan Financier Excel")
        st.caption("5 ans en Ariary — Cash Flow — SWOT — Point mort — Bilan prévisionnel")
        try:
            xlsx_bytes = generate_excel(d)
            st.download_button(
                label="⬇️ Télécharger le Plan Financier (.xlsx)",
                data=xlsx_bytes,
                file_name=f"PlanFinancier_{nom.replace(' ','_')}_Madagascar.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        except Exception as e:
            st.error(f"Erreur génération Excel : {e}")

    st.markdown("---")
    with st.expander("📋 Récapitulatif complet de vos réponses", expanded=False):
        cols = st.columns(2)
        items = [(k, v) for k, v in d.items() if v]
        half = len(items) // 2
        for i, (k, v) in enumerate(items):
            with cols[0 if i < half else 1]:
                st.markdown(f"**{k}** : {v}")

    if st.button("🔄 Créer un nouveau Business Plan"):
        st.session_state.step = 0
        st.session_state.data = {}
        st.session_state.ai_suggestions = {}
        st.rerun()


# ══════════════════════════════════════════════════════════════
# GÉNÉRATEUR WORD
# ══════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_colored_heading(doc, text, level=1, color="C8102E"):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_kv_table(doc, rows_data, col1_w=2.5, col2_w=4.2):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        lc = row.cells[0]
        lc.width = Inches(col1_w)
        set_cell_bg(lc, "FCE4EC" if i % 2 == 0 else "FFEBEE")
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9.5)
        lr.font.color.rgb = RGBColor.from_string("B71C1C")

        vc = row.cells[1]
        vc.width = Inches(col2_w)
        set_cell_bg(vc, "FFFFFF" if i % 2 == 0 else "FFF8F8")
        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value or "—"))
        vr.font.size = Pt(9.5)
    doc.add_paragraph()

def add_section_separator(doc, color="C8102E"):
    p = doc.add_paragraph()
    run = p.add_run("━" * 70)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(7)

def generate_word(d):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    nom = d.get("nom_entreprise", "Mon Projet") or "Mon Projet"

    # ── PAGE DE GARDE ───────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("🇲🇬  BUSINESS PLAN")
    r.font.size = Pt(34)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("C8102E")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(nom.upper())
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("1A1612")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(d.get("pitch", "") or "")
    r.font.size = Pt(12)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string("555555")

    doc.add_paragraph()
    doc.add_paragraph()
    add_kv_table(doc, [
        ("Porteur de projet", d.get("porteur", "—")),
        ("Région", d.get("region", "—")),
        ("Ville", d.get("ville", "—")),
        ("Secteur", d.get("domaine", "—")),
        ("Date de création prévue", d.get("date_creation", "—")),
        ("Statut juridique", d.get("statut_juridique", "—")),
        ("Format", "AFD / BNI Madagascar / Bailleurs Internationaux"),
        ("Date du document", datetime.now().strftime("%d/%m/%Y")),
    ])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document confidentiel — Business Plan Madagascar")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string("888888")

    doc.add_page_break()

    # ── SOMMAIRE ────────────────────────────────────────────────────────────
    add_colored_heading(doc, "TABLE DES MATIÈRES", 1, "C8102E")
    toc_items = [
        "1. Résumé Exécutif (Executive Summary)",
        "2. Présentation du Porteur de Projet",
        "3. L'Idée et la Vision",
        "4. Localisation et Contexte Madagascar",
        "5. Étude de Marché",
        "6. Produit / Service",
        "7. L'Équipe",
        "8. Plan de Financement",
        "9. Projections Financières sur 5 ans (en Ariary)",
        "10. Analyse SWOT",
        "11. Impact Social & Environnemental — ODD",
        "12. Analyse des Risques Pays",
        "13. Statut Juridique — Droit Malgache (OHADA)",
        "14. Conclusion et Appel à l'Action",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.3)
        p.runs[0].font.size = Pt(10.5)
        p.runs[0].font.color.rgb = RGBColor.from_string("374151")

    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────────────
    add_colored_heading(doc, "1. Résumé Exécutif — Executive Summary", 1, "C8102E")
    p = doc.add_paragraph(
        f"{nom} est un projet entrepreneurial basé à {d.get('ville','Madagascar')}, "
        f"dans la région {d.get('region','—')}, opérant dans le secteur "
        f"« {d.get('domaine','—')} »."
    )
    p.runs[0].font.size = Pt(11)
    if d.get("pitch"):
        doc.add_paragraph(d["pitch"]).runs[0].font.italic = True

    add_kv_table(doc, [
        ("Porteur", d.get("porteur", "—")),
        ("Secteur", d.get("domaine", "—")),
        ("Région d'activité", d.get("region", "—")),
        ("Cible commerciale", d.get("cible", "—")),
        ("Modèle économique", d.get("modele_eco", "—")),
        ("Statut juridique", d.get("statut_juridique", "—")),
        ("CA An 1 prévisionnel", fmt_ariary(d.get("ca_an1", 0))),
        ("CA An 3 prévisionnel", fmt_ariary(d.get("ca_an3", 0))),
        ("Budget de lancement", fmt_ariary(d.get("budget_lancement", 0))),
        ("Emplois créés An 1", f"{d.get('emplois_directs_an1', 0)} emplois directs"),
        ("Financement recherché", fmt_ariary(int(d.get("emprunt", 0) or 0) + int(d.get("montant_bailleur", 0) or 0))),
        ("Bailleur visé", d.get("bailleur_vise", "—")),
    ])
    doc.add_page_break()

    # ── 2. PORTEUR DE PROJET ─────────────────────────────────────────────
    add_colored_heading(doc, "2. Présentation du Porteur de Projet", 1, "C8102E")
    add_kv_table(doc, [
        ("Nom et prénom", d.get("porteur", "—")),
        ("Ville", d.get("ville", "—")),
        ("Région", d.get("region", "—")),
        ("Compétences", d.get("competences_porteur", "—")),
        ("Motivation", d.get("motivation", "—")),
    ])

    # ── 3. L'IDÉE ────────────────────────────────────────────────────────
    add_section_separator(doc)
    add_colored_heading(doc, "3. L'Idée et la Vision", 1, "C8102E")
    add_colored_heading(doc, "3.1 Problème identifié", 2, "007A3D")
    doc.add_paragraph(d.get("probleme", "Non renseigné"))
    add_colored_heading(doc, "3.2 Solution proposée", 2, "007A3D")
    doc.add_paragraph(d.get("solution", "Non renseigné"))
    add_colored_heading(doc, "3.3 Avantage Concurrentiel", 2, "007A3D")
    doc.add_paragraph(d.get("differenciateur", "Non renseigné"))
    doc.add_page_break()

    # ── 4. LOCALISATION ──────────────────────────────────────────────────
    add_colored_heading(doc, "4. Localisation et Contexte Madagascar", 1, "C8102E")
    add_kv_table(doc, [
        ("Région", d.get("region", "—")),
        ("Ville / Fokontany", d.get("ville", "—")),
        ("Zone d'activité", d.get("zone_activite", "—")),
        ("Rayon géographique clients", d.get("rayon_clients", "—")),
        ("Infrastructures disponibles", d.get("infrastructures", "—")),
        ("Contraintes / Opportunités région", d.get("contraintes_region", "—")),
    ])
    doc.add_page_break()

    # ── 5. ÉTUDE DE MARCHÉ ───────────────────────────────────────────────
    add_colored_heading(doc, "5. Étude de Marché", 1, "C8102E")
    add_kv_table(doc, [
        ("Taille du marché (Madagascar)", d.get("taille_marche", "—")),
        ("Source", d.get("source_marche", "—")),
        ("Maturité", d.get("maturite_marche", "—")),
        ("Tendances", d.get("tendances", "—")),
        ("Clients potentiels An 1", d.get("nb_clients_an1", "—")),
        ("Concurrents principaux", d.get("concurrents_details", "—")),
        ("Concurrents indirects", d.get("concurrents", "—")),
        ("Stratégie d'acquisition clients", d.get("strategie_acquisition", "—")),
    ])
    doc.add_page_break()

    # ── 6. PRODUIT / SERVICE ─────────────────────────────────────────────
    add_colored_heading(doc, "6. Produit / Service", 1, "C8102E")
    add_kv_table(doc, [
        ("Description", d.get("produit_detail", "—")),
        ("Production / Livraison", d.get("production", "—")),
        ("Modèle économique", d.get("modele_eco", "—")),
        ("Prix de vente", fmt_ariary(d.get("prix_vente_ar", 0))),
        ("Coût de revient", fmt_ariary(d.get("cout_revient_ar", 0))),
        ("Marge brute estimée", f"{d.get('marge_brute', 0)} %"),
        ("Ressources clés", d.get("ressources_cles", "—")),
        ("Fournisseurs principaux", d.get("fournisseurs", "—")),
    ])
    doc.add_page_break()

    # ── 7. ÉQUIPE ────────────────────────────────────────────────────────
    add_colored_heading(doc, "7. L'Équipe", 1, "C8102E")
    nb_str = d.get("seul_ou_plusieurs", "Seul(e)")
    nb = 1 if "Seul" in nb_str else (2 if "À 2" in nb_str else 3)
    for i in range(1, nb + 1):
        add_colored_heading(doc, f"Fondateur {i}", 2, "007A3D")
        add_kv_table(doc, [
            ("Nom", d.get(f"fond_{i}_nom", "—")),
            ("Rôle", d.get(f"fond_{i}_role", "—")),
            ("Formation", d.get(f"fond_{i}_formation", "—")),
            ("Expérience", d.get(f"fond_{i}_exp", "—")),
            ("Compétences", d.get(f"fond_{i}_skills", "—")),
        ])

    add_kv_table(doc, [
        ("Conseillers / Mentors", d.get("advisors", "—")),
        ("Recrutements prévus", d.get("recrutements", "—")),
    ])
    doc.add_page_break()

    # ── 8. FINANCEMENT ───────────────────────────────────────────────────
    add_colored_heading(doc, "8. Plan de Financement", 1, "C8102E")
    total_fin = (int(d.get("apports_perso", 0) or 0)
                 + int(d.get("love_money", 0) or 0)
                 + int(d.get("emprunt", 0) or 0)
                 + int(d.get("subventions", 0) or 0)
                 + int(d.get("montant_bailleur", 0) or 0))
    add_kv_table(doc, [
        ("Apports personnels", fmt_ariary(d.get("apports_perso", 0))),
        ("Apports famille / diaspora", fmt_ariary(d.get("love_money", 0))),
        ("Capital social", fmt_ariary(d.get("capital_social_ar", 0))),
        ("Emprunt bancaire", fmt_ariary(d.get("emprunt", 0))),
        ("Banque / MFI visée", d.get("banque_visee", "—")),
        ("Garanties disponibles", d.get("garanties", "—")),
        ("Subventions / Aides", fmt_ariary(d.get("subventions", 0))),
        ("Bailleur international", d.get("bailleur_vise", "—")),
        ("Montant bailleur", fmt_ariary(d.get("montant_bailleur", 0))),
        ("TOTAL FINANCEMENT", fmt_ariary(total_fin)),
        ("Budget de lancement", fmt_ariary(d.get("budget_lancement", 0))),
        ("Usage du financement", d.get("usage_budget", "—")),
    ])
    doc.add_page_break()

    # ── 9. PROJECTIONS FINANCIÈRES ────────────────────────────────────────
    add_colored_heading(doc, "9. Projections Financières — 5 ans (en Ariary)", 1, "C8102E")

    charges_fixes_mois = sum([
        int(d.get("charges_salaires", 0) or 0),
        int(d.get("charges_loyer", 0) or 0),
        int(d.get("charges_matieres", 0) or 0),
        int(d.get("charges_transport", 0) or 0),
        int(d.get("charges_utilities", 0) or 0),
        int(d.get("autres_charges", 0) or 0),
    ])
    cogs = int(d.get("cogs_pct", 35) or 35) / 100
    comm = int(d.get("commission_pct", 5) or 5) / 100

    proj_data = []
    for i in range(1, 6):
        ca = int(d.get(f"ca_an{i}", 0) or 0)
        charges_var = int(ca * (cogs + comm))
        charges_fix = charges_fixes_mois * 12
        resultat = ca - charges_var - charges_fix
        proj_data.append((f"Année {i}", ca, charges_var, charges_fix, resultat))

    table = doc.add_table(rows=1 + len(proj_data), cols=5)
    table.style = "Table Grid"
    headers = ["Année", "CA (Ar)", "Charges Variables (Ar)", "Charges Fixes (Ar)", "Résultat Net (Ar)"]
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        set_cell_bg(cell, "C8102E")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        r.font.size = Pt(9)

    for ri, (an, ca, cv, cf, res) in enumerate(proj_data, 1):
        row = table.rows[ri]
        vals = [an, fmt_ariary(ca), fmt_ariary(cv), fmt_ariary(cf), fmt_ariary(res)]
        bg = "FFF5F5" if ri % 2 == 0 else "FFFFFF"
        for ci, v in enumerate(vals):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(9)
            if ci == 4:
                r.font.color.rgb = RGBColor.from_string("007A3D" if res >= 0 else "C8102E")
                r.font.bold = True
    doc.add_paragraph()
    doc.add_page_break()

    # ── 10. ANALYSE SWOT ─────────────────────────────────────────────────
    add_colored_heading(doc, "10. Analyse SWOT", 1, "C8102E")
    swot_table = doc.add_table(rows=2, cols=2)
    swot_table.style = "Table Grid"
    swot_data = [
        ("Forces (Strengths)", "007A3D", "FBF8FF",
         [d.get("differenciateur",""), d.get("competences_porteur",""),
          f"Ancrage local — {d.get('region','')}", d.get(f"fond_1_skills","")]),
        ("Opportunités (Opportunities)", "1565C0", "E3F2FD",
         [d.get("tendances",""), d.get("taille_marche",""),
          d.get("bailleur_vise",""), d.get("strategie_acquisition","")]),
        ("Faiblesses (Weaknesses)", "E65100", "FFF3E0",
         [f"Capital initial limité ({fmt_ariary(d.get('capital_social_ar',0))})",
          "Notoriété à construire", "Dépendance aux fournisseurs locaux",
          "Infrastructure logistique (routes)"]),
        ("Menaces (Threats)", "B71C1C", "FFEBEE",
         [d.get("risques_pays","Instabilité politique / climatique"),
          "Concurrence informelle", "Fluctuation du taux de change Ar/€",
          "Coupures d'électricité"]),
    ]
    positions = [(0,0), (0,1), (1,0), (1,1)]
    for (r, c), (title, color, bg, items) in zip(positions, swot_data):
        cell = swot_table.rows[r].cells[c]
        cell.width = Inches(3.3)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.size = Pt(10)
        for item in items:
            if item:
                bp = cell.add_paragraph(f"• {str(item)[:120]}")
                bp.runs[0].font.size = Pt(9)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 11. IMPACT SOCIAL ─────────────────────────────────────────────────
    add_colored_heading(doc, "11. Impact Social & Environnemental", 1, "C8102E")
    add_kv_table(doc, [
        ("Emplois directs créés — An 1", f"{d.get('emplois_directs_an1', 0)}"),
        ("Emplois directs créés — An 3", f"{d.get('emplois_directs_an3', 0)}"),
        ("Emplois indirects estimés", f"{d.get('emplois_indirects', 0)}"),
        ("Femmes bénéficiaires directes", f"{d.get('femmes_beneficiaires', 0)}"),
        ("Populations ciblées", d.get("populations_cibles", "—")),
        ("ODD ciblés", d.get("odds_selectionnes", "—")),
        ("Contribution aux ODD", d.get("contribution_odd", "—")),
        ("Impact environnemental", d.get("impact_env", "—")),
    ])
    doc.add_page_break()

    # ── 12. RISQUES PAYS ─────────────────────────────────────────────────
    add_colored_heading(doc, "12. Analyse des Risques Pays — Madagascar", 1, "C8102E")
    add_kv_table(doc, [
        ("Risques identifiés", d.get("risques_pays", "—")),
        ("Mesures d'atténuation", d.get("attenuation_risques", "—")),
    ])
    doc.add_page_break()

    # ── 13. STATUT JURIDIQUE ─────────────────────────────────────────────
    add_colored_heading(doc, "13. Statut Juridique — Droit Malgache (OHADA)", 1, "C8102E")
    add_kv_table(doc, [
        ("Forme juridique", d.get("statut_juridique", "—")),
        ("Capital social", fmt_ariary(d.get("capital_social_ar", 0))),
        ("Siège social", d.get("siege_social", "—")),
        ("Immatriculation via", d.get("immatriculation_via", "—")),
        ("Régime fiscal", d.get("regime_fiscal", "—")),
    ])
    doc.add_page_break()

    # ── 14. CONCLUSION ───────────────────────────────────────────────────
    add_colored_heading(doc, "14. Conclusion et Appel à l'Action", 1, "C8102E")
    p = doc.add_paragraph(
        f"{nom} est un projet entrepreneurial à fort potentiel, ancré dans la réalité "
        f"économique et sociale de Madagascar. Avec un chiffre d'affaires prévisionnel de "
        f"{fmt_ariary(d.get('ca_an1',0))} dès la première année, {d.get('emplois_directs_an1',0)} "
        f"emplois directs créés, et un impact social concret sur les populations de "
        f"{d.get('region','la région')}, ce projet répond aux critères de financement "
        f"des principaux bailleurs opérant à Madagascar."
    )
    p.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph("Nous vous invitons à nous rejoindre dans cette aventure entrepreneuriale malgache.")
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor.from_string("007A3D")

    doc.add_paragraph()
    p = doc.add_paragraph(f"Porteur : {d.get('porteur','—')}  |  Date : {datetime.now().strftime('%d/%m/%Y')}")
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor.from_string("888888")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════
# GÉNÉRATEUR EXCEL
# ══════════════════════════════════════════════════════════════

def thin_border():
    s = Side(style="thin", color="DDDDDD")
    return Border(left=s, right=s, top=s, bottom=s)

def thick_border():
    s = Side(style="medium", color="C8102E")
    return Border(left=s, right=s, top=s, bottom=s)

def cell_style(cell, value, bold=False, color="000000", bg=None, fmt=None,
               align="left", size=10, border=True):
    cell.value = value
    cell.font = Font(name="Arial", bold=bold, size=size, color=color)
    if bg:
        cell.fill = PatternFill("solid", fgColor=bg)
    if fmt:
        cell.number_format = fmt
    cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
    if border:
        cell.border = thin_border()

def header_row(ws, row, headers, colors, widths=None):
    for ci, (h, col) in enumerate(zip(headers, colors), 1):
        c = ws.cell(row=row, column=ci, value=h)
        c.font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor=col)
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = thin_border()

def title_cell(ws, cell_ref, text, bg="C8102E", span=None):
    c = ws[cell_ref]
    c.value = text
    c.font = Font(name="Arial", bold=True, size=13, color="FFFFFF")
    c.fill = PatternFill("solid", fgColor=bg)
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[c.row].height = 36

def generate_excel(d):
    wb = openpyxl.Workbook()

    nom = d.get("nom_entreprise", "Projet") or "Projet"
    region = d.get("region", "Madagascar")
    porteur = d.get("porteur", "—")
    date_doc = datetime.now().strftime("%d/%m/%Y")

    # Données communes
    cas = [int(d.get(f"ca_an{i}", 0) or 0) for i in range(1, 6)]
    charges_salaires = int(d.get("charges_salaires", 0) or 0)
    charges_loyer    = int(d.get("charges_loyer", 0) or 0)
    charges_matieres = int(d.get("charges_matieres", 0) or 0)
    charges_transport= int(d.get("charges_transport", 0) or 0)
    charges_utilities= int(d.get("charges_utilities", 0) or 0)
    autres_charges   = int(d.get("autres_charges", 0) or 0)
    charges_fixes_mois = sum([charges_salaires, charges_loyer, charges_matieres,
                               charges_transport, charges_utilities, autres_charges])
    charges_fixes_an = charges_fixes_mois * 12
    cogs_pct  = int(d.get("cogs_pct", 35) or 35) / 100
    comm_pct  = int(d.get("commission_pct", 5) or 5) / 100
    invest_equipement  = int(d.get("invest_equipement", 0) or 0)
    invest_amenagement = int(d.get("invest_amenagement", 0) or 0)
    invest_fdr         = int(d.get("invest_fdr", 0) or 0)

    AR_FMT = '#,##0 "Ar"'

    # ── SHEET 1 : DASHBOARD KPIs ─────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "🇲🇬 Dashboard"
    ws1.column_dimensions["A"].width = 32
    ws1.column_dimensions["B"].width = 24
    ws1.column_dimensions["C"].width = 24
    ws1.column_dimensions["D"].width = 24

    ws1.merge_cells("A1:D1")
    title_cell(ws1, "A1", f"🇲🇬 BUSINESS PLAN — {nom.upper()} — {region}", "C8102E")

    ws1.merge_cells("A2:D2")
    c = ws1["A2"]
    c.value = f"Porteur : {porteur}  |  Date : {date_doc}  |  Format : AFD / BNI Madagascar"
    c.font = Font(name="Arial", size=10, color="FFFFFF", italic=True)
    c.fill = PatternFill("solid", fgColor="1A1612")
    c.alignment = Alignment(horizontal="center")
    ws1.row_dimensions[2].height = 22

    ws1.merge_cells("A3:D3")
    ws1["A3"].value = ""

    kpi_row = 4
    ws1.merge_cells(f"A{kpi_row}:D{kpi_row}")
    c = ws1.cell(kpi_row, 1, "📊 INDICATEURS CLÉS DU PROJET")
    c.font = Font(name="Arial", bold=True, size=11, color="C8102E")
    c.alignment = Alignment(horizontal="left")
    ws1.row_dimensions[kpi_row].height = 26

    kpis = [
        ("CA Année 1 (Ar)", cas[0], AR_FMT, "007A3D"),
        ("CA Année 3 (Ar)", cas[2], AR_FMT, "007A3D"),
        ("CA Année 5 (Ar)", cas[4], AR_FMT, "007A3D"),
        ("Charges Fixes/an (Ar)", charges_fixes_an, AR_FMT, "C8102E"),
        ("Résultat Net An 1 (Ar)", cas[0] - int(cas[0]*(cogs_pct+comm_pct)) - charges_fixes_an, AR_FMT, "1565C0"),
        ("Budget Lancement (Ar)", int(d.get("budget_lancement",0) or 0), AR_FMT, "E65100"),
        ("Emplois directs An 1", int(d.get("emplois_directs_an1",0) or 0), "#,##0", "007A3D"),
        ("Marge brute estimée (%)", int(d.get("marge_brute",0) or 0), "0%", "1565C0"),
    ]

    for ri, (label, val, fmt_str, col) in enumerate(kpis, kpi_row + 1):
        bg = "FFF5F5" if ri % 2 == 0 else "FFFFFF"
        c1 = ws1.cell(ri, 1, label)
        c1.font = Font(name="Arial", bold=True, size=10, color="374151")
        c1.fill = PatternFill("solid", fgColor=bg)
        c1.border = thin_border()

        c2 = ws1.cell(ri, 2, val)
        c2.number_format = fmt_str
        c2.font = Font(name="Arial", bold=True, size=11, color=col)
        c2.fill = PatternFill("solid", fgColor=bg)
        c2.border = thin_border()
        c2.alignment = Alignment(horizontal="right")

        ws1.row_dimensions[ri].height = 22

    # ── SHEET 2 : COMPTE DE RÉSULTAT 5 ANS ───────────────────────────────
    ws2 = wb.create_sheet("📈 Résultats 5 ans")
    ws2.column_dimensions["A"].width = 36
    for ci in range(2, 7):
        ws2.column_dimensions[get_column_letter(ci)].width = 22

    ws2.merge_cells("A1:F1")
    title_cell(ws2, "A1", f"COMPTE DE RÉSULTAT PRÉVISIONNEL — 5 ANS (Ar) — {nom.upper()}", "C8102E")

    header_row(ws2, 2,
               ["Poste", "Année 1", "Année 2", "Année 3", "Année 4", "Année 5"],
               ["1A1612","007A3D","007A3D","C8102E","C8102E","D4A017"])

    rows_data = []
    for i, ca in enumerate(cas, 1):
        cv = int(ca * (cogs_pct + comm_pct))
        cf = charges_fixes_an
        res = ca - cv - cf
        rows_data.append((f"Année {i}", ca, cv, cf, res))

    cr_items = [
        ("📥 Chiffre d'Affaires (CA)",         [r[1] for r in rows_data], "007A3D", True),
        ("📤 Charges Variables (COGS + comm.)", [r[2] for r in rows_data], "E65100", False),
        ("📤 Charges Fixes annuelles",          [r[3] for r in rows_data], "E65100", False),
        ("💹 Résultat Net",                     [r[4] for r in rows_data], "C8102E", True),
    ]

    for ri, (label, vals, col, bold) in enumerate(cr_items, 3):
        bg = "F8F8F8" if ri % 2 == 0 else "FFFFFF"
        c = ws2.cell(ri, 1, label)
        c.font = Font(name="Arial", bold=bold, size=10, color="374151")
        c.fill = PatternFill("solid", fgColor=bg)
        c.border = thin_border()
        for ci, v in enumerate(vals, 2):
            cell = ws2.cell(ri, ci, v)
            cell.number_format = AR_FMT
            cell.font = Font(name="Arial", bold=bold, size=10, color=col)
            cell.fill = PatternFill("solid", fgColor=bg)
            cell.border = thin_border()
            cell.alignment = Alignment(horizontal="right")
        ws2.row_dimensions[ri].height = 22

    # Graphique
    bar = BarChart()
    bar.title = "CA vs Résultat — 5 ans"
    bar.y_axis.title = "Ariary (Ar)"
    bar.width = 22; bar.height = 13
    data_ref = Reference(ws2, min_col=2, max_col=6, min_row=2, max_row=6)
    bar.add_data(data_ref, titles_from_data=True)
    ws2.add_chart(bar, "A10")

    # ── SHEET 3 : CASH FLOW ──────────────────────────────────────────────
    ws3 = wb.create_sheet("💸 Cash Flow 12 mois")
    ws3.column_dimensions["A"].width = 16
    for ci in range(2, 6):
        ws3.column_dimensions[get_column_letter(ci)].width = 22

    ws3.merge_cells("A1:E1")
    title_cell(ws3, "A1", f"PLAN DE TRÉSORERIE MENSUEL — ANNÉE 1 — {nom.upper()}", "007A3D")

    mois = ["Jan","Fév","Mar","Avr","Mai","Jun","Jul","Aoû","Sep","Oct","Nov","Déc"]
    header_row(ws3, 2,
               ["Mois","Encaissements (Ar)","Décaissements (Ar)","Solde Mensuel (Ar)","Solde Cumulé (Ar)"],
               ["1A1612","007A3D","C8102E","1565C0","D4A017"])

    ca1 = cas[0]
    ramp = [0.04,0.06,0.07,0.08,0.08,0.08,0.08,0.09,0.09,0.10,0.11,0.12]
    cumul = 0
    for ri, (m, r) in enumerate(zip(mois, ramp), 3):
        enc = int(ca1 * r)
        dec = int(enc * (cogs_pct + comm_pct)) + charges_fixes_mois
        solde = enc - dec
        cumul += solde
        bg = "F0FFF4" if solde >= 0 else "FFF0F0"

        c1 = ws3.cell(ri, 1, m)
        c1.font = Font(name="Arial", bold=True, size=10, color="374151")
        c1.fill = PatternFill("solid", fgColor=bg)
        c1.border = thin_border()

        for ci, (val, col) in enumerate([(enc,"007A3D"),(dec,"C8102E"),(solde,"1565C0"),(cumul,"D4A017")], 2):
            cell = ws3.cell(ri, ci, val)
            cell.number_format = AR_FMT
            cell.font = Font(name="Arial", size=10, color=col, bold=(ci >= 4))
            cell.fill = PatternFill("solid", fgColor=bg)
            cell.border = thin_border()
            cell.alignment = Alignment(horizontal="right")
        ws3.row_dimensions[ri].height = 20

    # ── SHEET 4 : FINANCEMENT ────────────────────────────────────────────
    ws4 = wb.create_sheet("💰 Financement")
    ws4.column_dimensions["A"].width = 38
    ws4.column_dimensions["B"].width = 24
    ws4.column_dimensions["C"].width = 18

    ws4.merge_cells("A1:C1")
    title_cell(ws4, "A1", "PLAN DE FINANCEMENT (Ariary)", "1A1612")

    header_row(ws4, 2, ["Source","Montant (Ar)","% Total"], ["C8102E","007A3D","D4A017"])

    fin_sources = [
        ("Apports personnels", int(d.get("apports_perso",0) or 0)),
        ("Apports famille / diaspora", int(d.get("love_money",0) or 0)),
        ("Emprunt bancaire", int(d.get("emprunt",0) or 0)),
        ("Subventions / Aides publiques", int(d.get("subventions",0) or 0)),
        ("Bailleur international", int(d.get("montant_bailleur",0) or 0)),
    ]
    total_fin = sum(v for _, v in fin_sources)

    for ri, (label, val) in enumerate(fin_sources, 3):
        pct = val / total_fin if total_fin > 0 else 0
        bg = "FFF5F5" if ri % 2 == 0 else "FFFFFF"

        c1 = ws4.cell(ri, 1, label)
        c1.font = Font(name="Arial", size=10, color="374151")
        c1.fill = PatternFill("solid", fgColor=bg)
        c1.border = thin_border()

        c2 = ws4.cell(ri, 2, val)
        c2.number_format = AR_FMT
        c2.font = Font(name="Arial", size=10, color="007A3D" if val > 0 else "AAAAAA")
        c2.fill = PatternFill("solid", fgColor=bg)
        c2.border = thin_border()
        c2.alignment = Alignment(horizontal="right")

        c3 = ws4.cell(ri, 3, pct)
        c3.number_format = "0.0%"
        c3.font = Font(name="Arial", size=10, color="C8102E")
        c3.fill = PatternFill("solid", fgColor=bg)
        c3.border = thin_border()
        c3.alignment = Alignment(horizontal="right")

    tr = len(fin_sources) + 3
    for ci, (val, fmt_str) in enumerate([("TOTAL",None),(total_fin,AR_FMT),(1.0,"0.0%")], 1):
        c = ws4.cell(tr, ci, val)
        c.font = Font(name="Arial", bold=True, size=11, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="1A1612")
        c.border = thick_border()
        if fmt_str:
            c.number_format = fmt_str
            c.alignment = Alignment(horizontal="right")

    # Budget vs financement
    budget = int(d.get("budget_lancement",0) or 0)
    ws4.cell(tr+2, 1, "Budget de lancement estimé").font = Font(name="Arial", bold=True, size=10, color="C8102E")
    ws4.cell(tr+2, 2, budget).number_format = AR_FMT
    ws4.cell(tr+2, 2).font = Font(name="Arial", bold=True, size=10, color="D4A017")
    ws4.cell(tr+2, 2).alignment = Alignment(horizontal="right")

    gap = total_fin - budget
    ws4.cell(tr+3, 1, "Solde (Financement - Budget)").font = Font(name="Arial", size=10, color="374151")
    ws4.cell(tr+3, 2, gap).number_format = AR_FMT
    ws4.cell(tr+3, 2).font = Font(name="Arial", bold=True, size=10, color="007A3D" if gap >= 0 else "C8102E")
    ws4.cell(tr+3, 2).alignment = Alignment(horizontal="right")

    # ── SHEET 5 : SEUIL DE RENTABILITÉ ───────────────────────────────────
    ws5 = wb.create_sheet("🎯 Seuil Rentabilité")
    ws5.column_dimensions["A"].width = 36
    ws5.column_dimensions["B"].width = 26

    ws5.merge_cells("A1:B1")
    title_cell(ws5, "A1", f"SEUIL DE RENTABILITÉ — POINT MORT (Ar) — {nom.upper()}", "C8102E")

    ca1_v = cas[0]
    taux_marge = 1 - cogs_pct - comm_pct
    point_mort = charges_fixes_an / taux_marge if taux_marge > 0 else 0
    mois_pm = (point_mort / ca1_v * 12) if ca1_v > 0 else 0

    pm_data = [
        ("Charges Fixes Annuelles (Ar)", charges_fixes_an, AR_FMT),
        ("Taux de Marge sur Coûts Variables (%)", taux_marge, "0.0%"),
        ("SEUIL DE RENTABILITÉ (Ar/an)", point_mort, AR_FMT),
        ("CA Annuel Prévisionnel An 1 (Ar)", ca1_v, AR_FMT),
        ("Mois pour atteindre le seuil (An 1)", round(mois_pm, 1), "#,##0.0"),
        ("Marge de sécurité (CA-Seuil) (Ar)", ca1_v - point_mort, AR_FMT),
        ("Ratio Marge de sécurité (%)", (ca1_v - point_mort) / ca1_v if ca1_v > 0 else 0, "0.0%"),
    ]

    header_row(ws5, 2, ["Indicateur", "Valeur"], ["C8102E", "007A3D"])
    for ri, (label, val, fmt_str) in enumerate(pm_data, 3):
        bg = "FFF5F5" if ri % 2 == 0 else "FFFFFF"
        bold = ri == 5  # seuil de rentabilité en gras

        c1 = ws5.cell(ri, 1, label)
        c1.font = Font(name="Arial", bold=bold, size=10, color="374151")
        c1.fill = PatternFill("solid", fgColor=bg)
        c1.border = thin_border()

        c2 = ws5.cell(ri, 2, val)
        c2.number_format = fmt_str
        c2.font = Font(name="Arial", bold=bold, size=10, color="C8102E" if bold else "374151")
        c2.fill = PatternFill("solid", fgColor=bg)
        c2.border = thin_border()
        c2.alignment = Alignment(horizontal="right")
        ws5.row_dimensions[ri].height = 22

    # ── SHEET 6 : IMPACT SOCIAL ──────────────────────────────────────────
    ws6 = wb.create_sheet("🌱 Impact Social")
    ws6.column_dimensions["A"].width = 38
    ws6.column_dimensions["B"].width = 24

    ws6.merge_cells("A1:B1")
    title_cell(ws6, "A1", f"IMPACT SOCIAL & ENVIRONNEMENTAL — {nom.upper()}", "007A3D")

    impact_data = [
        ("Emplois directs créés — Année 1", int(d.get("emplois_directs_an1",0) or 0), "#,##0"),
        ("Emplois directs créés — Année 3", int(d.get("emplois_directs_an3",0) or 0), "#,##0"),
        ("Emplois indirects estimés", int(d.get("emplois_indirects",0) or 0), "#,##0"),
        ("Femmes bénéficiaires directes", int(d.get("femmes_beneficiaires",0) or 0), "#,##0"),
        ("Populations ciblées", d.get("populations_cibles","—"), "@"),
        ("ODD ciblés", d.get("odds_selectionnes","—"), "@"),
        ("Contribution ODD", d.get("contribution_odd","—"), "@"),
        ("Risques identifiés", d.get("risques_pays","—"), "@"),
        ("Mesures d'atténuation", d.get("attenuation_risques","—"), "@"),
        ("Impact environnemental", d.get("impact_env","—"), "@"),
    ]

    header_row(ws6, 2, ["Indicateur", "Valeur / Description"], ["007A3D", "1565C0"])
    for ri, (label, val, fmt_str) in enumerate(impact_data, 3):
        bg = "F0FFF4" if ri % 2 == 0 else "FFFFFF"
        c1 = ws6.cell(ri, 1, label)
        c1.font = Font(name="Arial", bold=True, size=10, color="1B5E20")
        c1.fill = PatternFill("solid", fgColor=bg)
        c1.border = thin_border()

        c2 = ws6.cell(ri, 2, val)
        c2.number_format = fmt_str
        c2.font = Font(name="Arial", size=10, color="374151")
        c2.fill = PatternFill("solid", fgColor=bg)
        c2.border = thin_border()
        c2.alignment = Alignment(horizontal="left", wrap_text=True)
        ws6.row_dimensions[ri].height = 30

    # ── SHEET 7 : CHARGES DÉTAILLÉES ─────────────────────────────────────
    ws7 = wb.create_sheet("📉 Charges Détaillées")
    ws7.column_dimensions["A"].width = 34
    ws7.column_dimensions["B"].width = 22
    ws7.column_dimensions["C"].width = 22

    ws7.merge_cells("A1:C1")
    title_cell(ws7, "A1", f"DÉTAIL DES CHARGES (Ariary) — {nom.upper()}", "1A1612")

    header_row(ws7, 2, ["Poste de Charge", "Montant Mensuel (Ar)", "Montant Annuel (Ar)"],
               ["C8102E", "007A3D", "D4A017"])

    charges_list = [
        ("Salaires et rémunérations", charges_salaires),
        ("Loyer / hébergement", charges_loyer),
        ("Matières premières / Stock", charges_matieres),
        ("Transport / Carburant", charges_transport),
        ("Électricité / Eau / Téléphone", charges_utilities),
        ("Autres charges fixes", autres_charges),
    ]

    for ri, (label, val) in enumerate(charges_list, 3):
        bg = "FFF5F5" if ri % 2 == 0 else "FFFFFF"
        c1 = ws7.cell(ri, 1, label)
        c1.font = Font(name="Arial", size=10, color="374151")
        c1.fill = PatternFill("solid", fgColor=bg)
        c1.border = thin_border()

        c2 = ws7.cell(ri, 2, val)
        c2.number_format = AR_FMT
        c2.font = Font(name="Arial", size=10, color="374151")
        c2.fill = PatternFill("solid", fgColor=bg)
        c2.border = thin_border()
        c2.alignment = Alignment(horizontal="right")

        c3 = ws7.cell(ri, 3, val * 12)
        c3.number_format = AR_FMT
        c3.font = Font(name="Arial", size=10, color="C8102E")
        c3.fill = PatternFill("solid", fgColor=bg)
        c3.border = thin_border()
        c3.alignment = Alignment(horizontal="right")
        ws7.row_dimensions[ri].height = 20

    tr = len(charges_list) + 3
    for ci, (v, fmt_str) in enumerate([("TOTAL MENSUEL",None),(charges_fixes_mois,AR_FMT),(charges_fixes_an,AR_FMT)], 1):
        c = ws7.cell(tr, ci, v)
        c.font = Font(name="Arial", bold=True, size=11, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="C8102E")
        if fmt_str:
            c.number_format = fmt_str
            c.alignment = Alignment(horizontal="right")
        c.border = thick_border()

    # ── SAVE ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════
# ROUTER PRINCIPAL
# ══════════════════════════════════════════════════════════════
def main():
    step = st.session_state.step

    if step > 0:
        render_progress()

    steps_map = {
        0:  step_0_welcome,
        1:  step_1_idea,
        2:  step_2_project,
        3:  step_3_location,
        4:  step_4_market,
        5:  step_5_product,
        6:  step_6_team,
        7:  step_7_financing,
        8:  step_8_projections,
        9:  step_9_impact,
        10: step_10_legal,
        11: step_11_recap,
    }

    fn = steps_map.get(step, step_0_welcome)
    fn()


if __name__ == "__main__":
    main()
